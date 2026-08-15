#!/usr/bin/env python3
"""Unit tests for the /doc-text endpoint of hermes-bridge (no network calls).

Covers:
  (а) no/missing X-BRIDGE-TOKEN -> 401
  (б) bad request (not JSON)    -> 400
  (в) unsupported type (.zip)   -> 400
  (г) pypdf missing             -> 503 (mocked import)
  (д) extract_text for .txt     -> real file, full handler flow
  (е) digest with mocked hermes CLI -> <TEXT>...</TEXT> marker parsing + fallbacks

Extra: TELEGRAM_BOT_TOKEN empty -> 503; Telegram network errors -> 400/502;
getFile ok:false -> 400; digest=false skips the LLM call; docx/pdf unit tests
skip unless the library is installed.

Run: python3 test_doc_text.py
"""
import importlib.util
import io
import json
import os
import subprocess
import sys
import tempfile
import unittest
from contextlib import ExitStack
from unittest import mock

DIR = os.path.dirname(os.path.abspath(__file__))
SERVER_PATH = os.path.join(DIR, "hermes-bridge-server.py")

spec = importlib.util.spec_from_file_location("hermes_bridge_server", SERVER_PATH)
hbserver = importlib.util.module_from_spec(spec)
spec.loader.exec_module(hbserver)


class _Body(object):
    def __init__(self, data):
        self.data = data

    def read(self, n=-1):
        return self.data


class FakeHandler(hbserver.Handler):
    """Handler with socket I/O replaced by in-memory objects."""

    def __init__(self, body=b"{}", headers=None):
        self.rfile = _Body(body)
        self.headers = headers or {}
        self.sent = None

    def _send(self, code, obj):
        self.sent = (code, obj)


def make_handler(body=None, token="secret", extra_headers=None, path="/doc-text"):
    body = json.dumps(body).encode("utf-8") if not isinstance(body, bytes) else body
    headers = {"X-BRIDGE-TOKEN": token, "Content-Length": str(len(body))}
    headers.update(extra_headers or {})
    h = FakeHandler(body, headers)
    h.path = path
    return h


class DocTextValidationTests(unittest.TestCase):
    """(а) auth, (б) bad request, (в) unsupported type, telegram not configured."""

    def setUp(self):
        self._old_token = hbserver.TOKEN
        hbserver.TOKEN = "secret"

    def tearDown(self):
        hbserver.TOKEN = self._old_token

    def test_no_token_401(self):
        hbserver.TOKEN = ""
        h = make_handler({"file_id": "f", "file_name": "a.txt"})
        h.do_POST()
        self.assertEqual(h.sent[0], 401)

    def test_wrong_token_401(self):
        h = make_handler({"file_id": "f", "file_name": "a.txt"}, token="nope")
        h.do_POST()
        self.assertEqual(h.sent[0], 401)

    def test_bad_json_400(self):
        h = make_handler(b"this is not json")
        h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertIn("bad request", h.sent[1]["error"])

    def test_missing_file_id_400(self):
        h = make_handler({"file_name": "a.txt"})
        h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertIn("file_id", h.sent[1]["error"])

    def test_missing_file_name_400(self):
        h = make_handler({"file_id": "f"})
        h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertIn("file_name", h.sent[1]["error"])

    def test_unsupported_type_400(self):
        # must not touch the network: ext check happens before Telegram calls
        h = make_handler({"file_id": "f", "file_name": "x.zip"})
        h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertEqual(h.sent[1]["error"], "unsupported type")

    def test_unsupported_type_no_ext_400(self):
        h = make_handler({"file_id": "f", "file_name": "README"})
        h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertEqual(h.sent[1]["error"], "unsupported type")

    def test_telegram_not_configured_503(self):
        with mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": ""}):
            h = make_handler({"file_id": "f", "file_name": "a.txt"})
            h.do_POST()
        self.assertEqual(h.sent[0], 503)
        self.assertEqual(h.sent[1]["error"], "telegram not configured")

    def test_unknown_path_404(self):
        h = make_handler({"file_id": "f", "file_name": "a.txt"}, path="/nope")
        h.do_POST()
        self.assertEqual(h.sent[0], 404)

    def test_health_unchanged(self):
        h = FakeHandler(b"", {"X-BRIDGE-TOKEN": "secret"})
        h.path = "/health"
        h.do_GET()
        self.assertEqual(h.sent, (200, {"ok": True}))


class ExtractTests(unittest.TestCase):
    """(г) missing libs -> 503; (д) real-file txt extraction + full flow."""

    def setUp(self):
        self._old_token = hbserver.TOKEN
        hbserver.TOKEN = "secret"
        self.tmpdir = tempfile.mkdtemp(prefix="doc-text-test-")

    def tearDown(self):
        hbserver.TOKEN = self._old_token

    def _fixture(self, name, content):
        path = os.path.join(self.tmpdir, name)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return path

    def test_pypdf_missing_503(self):
        real_import = __import__

        def fake_import(name, *args, **kwargs):
            if name == "pypdf":
                raise ImportError("No module named 'pypdf'")
            return real_import(name, *args, **kwargs)

        fixture = self._fixture("doc.pdf", "not a real pdf, import must fail first")
        with ExitStack() as st:
            st.enter_context(mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}))
            st.enter_context(mock.patch.object(hbserver, "_telegram_get_file",
                                               return_value="docs/doc.pdf"))
            st.enter_context(mock.patch.object(hbserver, "_telegram_download"))
            st.enter_context(mock.patch.object(hbserver, "_tmp_path", return_value=fixture))
            st.enter_context(mock.patch("builtins.__import__", side_effect=fake_import))
            h = make_handler({"file_id": "f", "file_name": "doc.pdf"})
            h.do_POST()
        self.assertEqual(h.sent[0], 503)
        self.assertIn("pypdf missing", h.sent[1]["error"])

    def test_python_docx_missing_503(self):
        real_import = __import__

        def fake_import(name, *args, **kwargs):
            if name == "docx":
                raise ImportError("No module named 'docx'")
            return real_import(name, *args, **kwargs)

        fixture = self._fixture("doc.docx", "not a real docx")
        with ExitStack() as st:
            st.enter_context(mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}))
            st.enter_context(mock.patch.object(hbserver, "_telegram_get_file",
                                               return_value="docs/doc.docx"))
            st.enter_context(mock.patch.object(hbserver, "_telegram_download"))
            st.enter_context(mock.patch.object(hbserver, "_tmp_path", return_value=fixture))
            st.enter_context(mock.patch("builtins.__import__", side_effect=fake_import))
            h = make_handler({"file_id": "f", "file_name": "doc.docx"})
            h.do_POST()
        self.assertEqual(h.sent[0], 503)
        self.assertIn("python-docx missing", h.sent[1]["error"])

    def test_extract_text_txt_real_file(self):
        path = self._fixture("sample.txt", "Привет, контент-завод!\nвторая строка")
        self.assertEqual(hbserver._extract_text(path, "txt"),
                         "Привет, контент-завод!\nвторая строка")

    def test_extract_text_md_real_file(self):
        path = self._fixture("sample.md", "# Заголовок\n\nтело")
        self.assertEqual(hbserver._extract_text(path, "md"), "# Заголовок\n\nтело")

    @unittest.skipUnless(importlib.util.find_spec("docx"), "python-docx not installed")
    def test_extract_text_docx_real_file(self):
        import docx
        path = os.path.join(self.tmpdir, "sample.docx")
        d = docx.Document()
        d.add_paragraph("Первый абзац")
        d.add_paragraph("Второй абзац")
        d.save(path)
        text = hbserver._extract_text(path, "docx")
        self.assertIn("Первый абзац", text)
        self.assertIn("Второй абзац", text)

    @unittest.skipUnless(importlib.util.find_spec("pypdf"), "pypdf not installed")
    def test_extract_text_pdf_real_file(self):
        from pypdf import PdfReader, PdfWriter
        path = os.path.join(self.tmpdir, "sample.pdf")
        # minimal one-page PDF with a text object
        content = (
            "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj "
            "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj "
            "3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] "
            "/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj "
            "4 0 obj << /Length 60 >> stream\n"
            "BT /F1 12 Tf 20 100 Td (Hello pypdf) Tj ET\n"
            "endstream endobj "
            "5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj "
            "trailer << /Root 1 0 R >>\n%%EOF"
        )
        with open(path, "w") as f:
            f.write(content)
        text = hbserver._extract_text(path, "pdf")
        self.assertIn("Hello pypdf", text)

    def test_txt_full_flow_truncation(self):
        content = ""
        i = 0
        while len(content) < 31000:
            content += "строка номер %05d\n" % i
            i += 1
        fixture = self._fixture("sample.txt", content)
        with ExitStack() as st:
            st.enter_context(mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}))
            st.enter_context(mock.patch.object(hbserver, "_telegram_get_file",
                                               return_value="docs/sample.txt"))
            st.enter_context(mock.patch.object(hbserver, "_telegram_download"))
            st.enter_context(mock.patch.object(hbserver, "_tmp_path", return_value=fixture))
            h = make_handler({"file_id": "f", "file_name": "sample.txt"})
            h.do_POST()
        self.assertEqual(h.sent[0], 200)
        body = h.sent[1]
        self.assertEqual(body["name"], "sample.txt")
        self.assertEqual(body["mime"], "text/plain")
        self.assertEqual(body["chars"], 30000)
        self.assertEqual(len(body["text"]), 30000)
        self.assertFalse(body["digested"])
        self.assertTrue(body["text"].startswith("строка номер 00000"))

    def test_tmp_file_removed_after_flow(self):
        fixture = self._fixture("sample.txt", "короткий текст")
        with ExitStack() as st:
            st.enter_context(mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}))
            st.enter_context(mock.patch.object(hbserver, "_telegram_get_file",
                                               return_value="docs/sample.txt"))
            st.enter_context(mock.patch.object(hbserver, "_telegram_download"))
            st.enter_context(mock.patch.object(hbserver, "_tmp_path", return_value=fixture))
            h = make_handler({"file_id": "f", "file_name": "sample.txt"})
            h.do_POST()
        self.assertEqual(h.sent[0], 200)
        self.assertFalse(os.path.exists(fixture))


class TelegramErrorTests(unittest.TestCase):
    """getFile/download failure mapping (400/502), no network."""

    def setUp(self):
        self._old_token = hbserver.TOKEN
        hbserver.TOKEN = "secret"

    def tearDown(self):
        hbserver.TOKEN = self._old_token

    def test_getfile_ok_false_400(self):
        resp = mock.MagicMock()
        resp.read.return_value = json.dumps(
            {"ok": False, "error_code": 400,
             "description": "Bad Request: file is too big"}
        ).encode("utf-8")
        resp.__enter__.return_value = resp  # urlopen(...) used as context manager
        with mock.patch.object(hbserver.urllib.request, "urlopen", return_value=resp):
            with self.assertRaises(hbserver.TelegramError) as cm:
                hbserver._telegram_get_file("tok", "fid")
        self.assertEqual(cm.exception.code, 400)
        self.assertIn("file is too big", str(cm.exception))

    def test_getfile_http_error_400(self):
        err = __import__("urllib.error").error.HTTPError(
            "https://api.telegram.org/botT/getFile", 400, "Bad Request",
            {}, io.BytesIO(b""))
        with mock.patch.object(hbserver.urllib.request, "urlopen", side_effect=err):
            with self.assertRaises(hbserver.TelegramError) as cm:
                hbserver._telegram_get_file("tok", "fid")
        self.assertEqual(cm.exception.code, 400)

    def test_handler_network_error_502(self):
        with ExitStack() as st:
            st.enter_context(mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}))
            st.enter_context(mock.patch.object(
                hbserver, "_telegram_get_file",
                side_effect=hbserver.TelegramError("telegram unreachable: test", 502)))
            h = make_handler({"file_id": "f", "file_name": "a.txt"})
            h.do_POST()
        self.assertEqual(h.sent[0], 502)
        self.assertIn("telegram unreachable", h.sent[1]["error"])

    def test_handler_file_not_found_400(self):
        with ExitStack() as st:
            st.enter_context(mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}))
            st.enter_context(mock.patch.object(
                hbserver, "_telegram_get_file",
                side_effect=hbserver.TelegramError("telegram file not found (HTTP 404)", 400)))
            h = make_handler({"file_id": "f", "file_name": "a.txt"})
            h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertIn("not found", h.sent[1]["error"])


class DigestTests(unittest.TestCase):
    """(е) digest via mocked hermes CLI: markers, fallbacks, digest=false."""

    def setUp(self):
        self._old_token = hbserver.TOKEN
        hbserver.TOKEN = "secret"
        self.tmpdir = tempfile.mkdtemp(prefix="doc-text-digest-")

    def tearDown(self):
        hbserver.TOKEN = self._old_token

    def _base_mocks(self, txt="короткий текст документа для дайджеста"):
        path = os.path.join(self.tmpdir, "sample.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write(txt)
        return [
            mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "tok"}),
            mock.patch.object(hbserver, "_telegram_get_file", return_value="docs/sample.txt"),
            mock.patch.object(hbserver, "_telegram_download"),
            mock.patch.object(hbserver, "_tmp_path", return_value=path),
        ]

    def _run(self, body, txt="короткий текст документа для дайджеста"):
        with ExitStack() as st:
            for p in self._base_mocks(txt):
                st.enter_context(p)
            run_patch = mock.patch.object(hbserver.subprocess, "run", return_value=self.fake)
            run_mock = run_patch.start()
            st.callback(run_patch.stop)
            h = make_handler(body)
            h.do_POST()
        return h.sent, run_mock

    def test_digest_marker_parsing(self):
        self.fake = mock.Mock()
        self.fake.returncode = 0
        self.fake.stderr = ""
        self.fake.stdout = (
            "рассуждения перед ответом\n"
            "<TEXT>Это краткий дайджест документа в маркерах.</TEXT>\n"
            "хвост после маркера"
        )
        sent, _ = self._run({"file_id": "f", "file_name": "sample.txt", "digest": True})
        self.assertEqual(sent[0], 200)
        body = sent[1]
        self.assertTrue(body["digested"])
        self.assertEqual(body["text"], "Это краткий дайджест документа в маркерах.")
        self.assertEqual(body["chars"], len(body["text"]))
        self.assertEqual(body["name"], "sample.txt")
        self.assertEqual(body["mime"], "text/plain")

    def test_digest_uses_hermes_cli_flags(self):
        self.fake = mock.Mock()
        self.fake.returncode = 0
        self.fake.stderr = ""
        self.fake.stdout = "<TEXT>дайджест</TEXT>"
        sent, run = self._run({"file_id": "f", "file_name": "sample.txt", "digest": True})
        self.assertEqual(sent[0], 200)
        cmd = run.call_args[0][0]
        self.assertEqual(cmd[0], hbserver.HERMES_BIN)
        self.assertEqual(cmd[1], "chat")
        self.assertIn("-q", cmd)
        self.assertIn("--cli", cmd)
        self.assertIn("-Q", cmd)
        self.assertIn("--reasoning", cmd)
        self.assertIn(hbserver.DIGEST_PROMPT, cmd[3])
        self.assertIn("короткий текст документа", cmd[3])

    def test_digest_no_markers_fallback(self):
        txt = "длинный текст " * 500  # ~7500 chars > 2000 fallback limit
        self.fake = mock.Mock()
        self.fake.returncode = 0
        self.fake.stderr = ""
        self.fake.stdout = "ответ без TEXT-маркеров"
        sent, _ = self._run({"file_id": "f", "file_name": "sample.txt", "digest": True},
                            txt=txt)
        self.assertEqual(sent[0], 200)
        body = sent[1]
        self.assertFalse(body["digested"])
        self.assertEqual(body["chars"], 2000)
        self.assertEqual(body["text"], txt[:2000])

    def test_digest_timeout_fallback(self):
        self.fake = mock.Mock(side_effect=subprocess.TimeoutExpired(
            ["hermes", "chat"], timeout=hbserver.DIGEST_TIMEOUT))
        sent, _ = self._run({"file_id": "f", "file_name": "sample.txt", "digest": True})
        self.assertEqual(sent[0], 200)
        body = sent[1]
        self.assertFalse(body["digested"])
        self.assertEqual(body["text"], "короткий текст документа для дайджеста")
        self.assertEqual(body["chars"], len(body["text"]))

    def test_digest_false_skips_llm(self):
        self.fake = mock.Mock()
        sent, run = self._run({"file_id": "f", "file_name": "sample.txt", "digest": False})
        self.assertEqual(sent[0], 200)
        self.assertFalse(sent[1]["digested"])
        run.assert_not_called()

    def test_digest_string_true_parsed(self):
        self.fake = mock.Mock()
        self.fake.returncode = 0
        self.fake.stderr = ""
        self.fake.stdout = "<TEXT>дайджест из строкового true</TEXT>"
        sent, _ = self._run({"file_id": "f", "file_name": "sample.txt", "digest": "true"})
        self.assertEqual(sent[0], 200)
        self.assertTrue(sent[1]["digested"])


class AskRegressionTests(unittest.TestCase):
    """/ask behaviour must stay unchanged."""

    def setUp(self):
        self._old_token = hbserver.TOKEN
        hbserver.TOKEN = "secret"

    def tearDown(self):
        hbserver.TOKEN = self._old_token

    def test_ask_still_works(self):
        fake = mock.Mock()
        fake.returncode = 0
        fake.stdout = "ответ"
        fake.stderr = ""
        with mock.patch.object(hbserver.subprocess, "run", return_value=fake) as run:
            h = make_handler({"skill": "analyst", "prompt": "вопрос"},
                             path="/ask")
            h.do_POST()
        self.assertEqual(h.sent[0], 200)
        self.assertEqual(h.sent[1]["answer"], "ответ")
        cmd = run.call_args[0][0]
        self.assertIn("-s", cmd)
        self.assertIn("content-factory/analyst", cmd)

    def test_ask_invalid_skill_400(self):
        h = make_handler({"skill": "nope", "prompt": "вопрос"}, path="/ask")
        h.do_POST()
        self.assertEqual(h.sent[0], 400)
        self.assertIn("invalid skill", h.sent[1]["error"])


if __name__ == "__main__":
    unittest.main(verbosity=2)
